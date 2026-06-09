"""
Phase 3 — Export a generated application package to .docx / .pdf / .md.

A package is a plain dict: ``{"title": str, "sections": [{"title", "content"}]}``
(as produced by DocumentService / stored on ApplicationPackage). All three
renderers consume that same shape, so the document model is defined once.

- Markdown is pure-Python and always available (also used as the test oracle).
- .docx uses ``python-docx`` and .pdf uses ``reportlab`` — both imported lazily
  and reported via a clear RuntimeError when missing, so the dependency set
  stays optional and the rest of the backend imports without them.
"""
from __future__ import annotations

import io
from dataclasses import dataclass

# (key, mime, filename-extension) for each supported format.
EXPORT_FORMATS = {
    "md": ("text/markdown", "md"),
    "docx": (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "docx",
    ),
    "pdf": ("application/pdf", "pdf"),
}


@dataclass
class RenderedDocument:
    content: bytes
    media_type: str
    extension: str


def _sections(package: dict) -> list[dict]:
    return package.get("sections") or []


def _safe_filename(title: str) -> str:
    """Slug a package title into an ASCII filename stem."""
    keep = [c if c.isalnum() or c in " -_" else "_" for c in (title or "application")]
    stem = "".join(keep).strip().replace(" ", "_")[:80]
    return stem or "application"


def to_markdown(package: dict) -> bytes:
    """Render the package as Markdown (always available)."""
    lines = [f"# {package.get('title', 'Grant Application')}", ""]
    for sec in _sections(package):
        lines.append(f"## {sec.get('title', '')}".rstrip())
        lines.append("")
        lines.append((sec.get("content") or "").strip())
        lines.append("")
    return "\n".join(lines).encode("utf-8")


def to_docx(package: dict) -> bytes:
    """Render the package as a .docx (requires python-docx)."""
    try:
        from docx import Document
    except ImportError as e:  # pragma: no cover - exercised only without the dep
        raise RuntimeError(
            "DOCX export requires python-docx. Install it with "
            "`pip install python-docx`."
        ) from e

    doc = Document()
    doc.add_heading(package.get("title", "Grant Application"), level=0)
    for sec in _sections(package):
        doc.add_heading(sec.get("title", ""), level=1)
        for para in (sec.get("content") or "").split("\n"):
            doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def to_pdf(package: dict) -> bytes:
    """Render the package as a .pdf (requires reportlab)."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
        )
        from reportlab.lib.units import cm
    except ImportError as e:  # pragma: no cover - exercised only without the dep
        raise RuntimeError(
            "PDF export requires reportlab. Install it with `pip install reportlab`."
        ) from e

    from xml.sax.saxutils import escape

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, title=package.get("title", ""))
    styles = getSampleStyleSheet()
    story = [Paragraph(escape(package.get("title", "Grant Application")), styles["Title"])]
    story.append(Spacer(1, 0.5 * cm))
    for sec in _sections(package):
        story.append(Paragraph(escape(sec.get("title", "")), styles["Heading1"]))
        for para in (sec.get("content") or "").split("\n"):
            if para.strip():
                story.append(Paragraph(escape(para), styles["BodyText"]))
            else:
                story.append(Spacer(1, 0.2 * cm))
        story.append(Spacer(1, 0.3 * cm))
    doc.build(story)
    return buf.getvalue()


def render(package: dict, fmt: str) -> RenderedDocument:
    """Render ``package`` to the requested format key ('md'|'docx'|'pdf')."""
    if fmt not in EXPORT_FORMATS:
        raise ValueError(
            f"Unsupported export format '{fmt}'. Choose one of: "
            f"{', '.join(EXPORT_FORMATS)}"
        )
    media_type, ext = EXPORT_FORMATS[fmt]
    renderer = {"md": to_markdown, "docx": to_docx, "pdf": to_pdf}[fmt]
    return RenderedDocument(content=renderer(package), media_type=media_type, extension=ext)
